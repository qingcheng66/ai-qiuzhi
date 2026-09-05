"""导出服务：HTML / PDF(WeasyPrint) / DOCX(python-docx)

模板机制：
- 内置模板：Jinja2 模板文件 (app/template_static/*.j2)，接收 resume content(dict)
- 自定义/导入模板：数据库中 content 字段为 Jinja2 模板 HTML，接收扁平化变量 + body
"""
import io
from pathlib import Path
from typing import Any

import markdown as md_lib
from docx import Document as DocxDocument
from jinja2 import Environment
from weasyprint import HTML as WeasyHTML

from app.models.template import Template

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "template_static"

BUILTIN_TEMPLATES: dict[str, str] = {}

_env = Environment(autoescape=True)
_env.filters["md"] = lambda t: md_lib.markdown(t or "", extensions=["extra", "nl2br"])


def load_builtin_templates() -> None:
    """加载内置模板文件到模块字典"""
    global BUILTIN_TEMPLATES
    BUILTIN_TEMPLATES = {}
    for f in TEMPLATE_DIR.glob("*.j2"):
        BUILTIN_TEMPLATES[f.stem] = f.read_text(encoding="utf-8")
    if not BUILTIN_TEMPLATES:
        BUILTIN_TEMPLATES["minimal"] = _default_template()


def _default_template() -> str:
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:'PingFang SC',sans-serif;padding:40px 48px;color:#1f2937;"
        "line-height:1.6} h1{color:#111827} h2{color:#2563eb;border-left:4px solid #2563eb;"
        "padding-left:8px} .item{margin-bottom:10px}.item h3{margin:0}.contact{color:#4b5563}"
        "ul{padding-left:18px}</style></head><body>{{ body }}</body></html>"
    )


def _flatten(data: dict[str, Any]) -> dict[str, Any]:
    """把 resume content 扁平化为模板变量 + 附加 body 供后备渲染"""
    basics = data.get("basics", {}) or {}
    custom_fields = basics.get("custom_fields", []) or []
    
    # 抽取常规字段
    flat = {
        "name": basics.get("name", "姓名"),
        "label": basics.get("label", ""),
        "email": basics.get("email", ""),
        "phone": basics.get("phone", ""),
        "location": basics.get("location", ""),
        "github": basics.get("github", ""),
        "blog": basics.get("blog", ""),
        "summary": basics.get("summary", ""),
        "birthDate": basics.get("birthDate") or basics.get("birth", ""),
        "photo": basics.get("photo") or basics.get("avatar", ""),
        "education": data.get("education", []),
        "skills": data.get("skills", []),
        "projects": data.get("projects", []),
        "experience": data.get("experience", []) or data.get("work", []),
        "highlights": data.get("highlights", []),
        "custom_sections": data.get("custom_sections", []),
        "custom_fields": custom_fields,
        "meta": data.get("meta", {}),
        "data": data,
        "body": _body_html(data),
    }

    # 动态把自定义板块和字段注入到平铺命名空间中
    for sec in data.get("custom_sections", []) or []:
        sec_id = sec.get("id")
        if sec_id:
            flat[sec_id] = sec.get("items", []) or sec.get("content", "")

    for cf in custom_fields:
        k = cf.get("key") or cf.get("label")
        if k:
            flat[k] = cf.get("value", "")

    return flat


def _block(title: str, rows: list[dict]) -> str:
    if not rows:
        return ""
    parts = [f'<section class="block"><h2>{title}</h2>']
    for r in rows or []:
        head = " · ".join(str(r.get(k, "")) for k in ("name", "company", "institution", "role", "title") if r.get(k))
        parts.append(f'<div class="item"><h3>{head}</h3>')
        date = " - ".join(str(x) for x in (r.get("startDate", ""), r.get("endDate", "")) if x)
        if date:
            parts.append(f'<span class="date">{date}</span>')
        for h in r.get("highlights", []) or []:
            parts.append(f"<li>{h}</li>")
        for kw in r.get("keywords", []) or []:
            parts.append(f"<li>{kw}</li>")
        parts.append("</div>")
    parts.append("</section>")
    return "".join(parts)


def _body_html(data: dict[str, Any]) -> str:
    basics = data.get("basics", {}) or {}
    contact_parts = [basics.get("email", ""), basics.get("phone", ""), basics.get("location", "")]
    for cf in basics.get("custom_fields", []) or []:
        if cf.get("label") and cf.get("value"):
            contact_parts.append(f"{cf['label']}: {cf['value']}")
    
    contacts_str = " · ".join(p for p in contact_parts if p)
    html = (
        f'<header><h1>{basics.get("name", "姓名")}</h1>'
        f'<p class="contact">{contacts_str}</p>'
        f'<p class="label">{basics.get("label", "")}</p></header>'
    )
    html += _block("教育背景", data.get("education", []))
    html += _block("专业技能", data.get("skills", []))
    html += _block("工作经历", data.get("experience", []) or data.get("work", []))
    html += _block("项目经历", data.get("projects", []))
    if data.get("highlights"):
        rows = [
            {"name": h.get("title", ""), "highlights": [h.get("content", "")]}
            for h in data["highlights"]
        ]
        html += _block("技术亮点", rows)

    # 渲染自定义板块
    for sec in data.get("custom_sections", []) or []:
        sec_title = sec.get("title") or "自定义模块"
        items = sec.get("items", [])
        if items:
            html += _block(sec_title, items)
        elif sec.get("content"):
            html += f'<section class="block"><h2>{sec_title}</h2><p>{sec["content"]}</p></section>'

    return html


def _builtin_html(name: str, data: dict[str, Any]) -> str:
    tpl = BUILTIN_TEMPLATES.get(name) or BUILTIN_TEMPLATES["minimal"]
    return _env.from_string(tpl).render(**_flatten(data))


def _custom_html(content: str, data: dict[str, Any]) -> str:
    return _env.from_string(content).render(**_flatten(data))


def render_html(content: dict[str, Any], template: Template | None = None) -> tuple[str, str]:
    """返回 (html, 使用的模板名)"""
    if template:
        html = _custom_html(template.content, content)
        return html, template.name
    name = (content.get("meta", {}).get("template") or "minimal")
    return _builtin_html(name, content), name


def export_pdf(content: dict[str, Any], template: Template | None = None) -> bytes:
    html, _ = render_html(content, template)
    return WeasyHTML(string=html).write_pdf()


def export_html(content: dict[str, Any], template: Template | None = None) -> str:
    html, _ = render_html(content, template)
    return html


def export_docx(content: dict[str, Any], template: Template | None = None) -> bytes:
    doc = DocxDocument()
    basics = content.get("basics", {}) or {}
    doc.add_heading(basics.get("name") or "简历", level=0)
    doc.add_paragraph(
        " · ".join(x for x in (basics.get("email"), basics.get("phone"), basics.get("location")) if x)
    )
    if basics.get("label"):
        doc.add_paragraph(basics["label"])

    def section(title, rows):
        if not rows:
            return
        doc.add_heading(title, level=1)
        for r in rows or []:
            head = " | ".join(str(r.get(k, "")) for k in ("name", "company", "institution", "role") if r.get(k))
            if head:
                p = doc.add_paragraph()
                r0 = p.add_run(head)
                r0.bold = True
            for h in r.get("highlights", []) or []:
                doc.add_paragraph("• " + str(h), style="List Bullet")
            for kw in r.get("keywords", []) or []:
                doc.add_paragraph("• " + str(kw), style="List Bullet")

    section("教育背景", content.get("education", []))
    section("专业技能", content.get("skills", []))
    section("工作经历", content.get("experience", []) or content.get("work", []))
    section("项目经历", content.get("projects", []))
    if content.get("highlights"):
        doc.add_heading("技术亮点", level=1)
        for h in content["highlights"]:
            doc.add_paragraph(h.get("title", ""), style="List Bullet")
            if h.get("content"):
                doc.add_paragraph(h["content"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()